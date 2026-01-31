from docx import Document
import re

REDACT = "████████"


def redact_word(input_path, output_path):
    doc = Document(input_path)

    for para in doc.paragraphs:
        text = para.text
        original = text

        # Email addresses
        text = re.sub(
            r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+",
            REDACT,
            text,
        )

        # Phone numbers (10 digits)
        text = re.sub(r"\b\d{10}\b", REDACT, text)

        # Long numbers (IDs)
        text = re.sub(r"\b\d{4,}\b", REDACT, text)

        # Sensitive keywords (case-insensitive)
        text = re.sub(
            r"\b(Aadhaar|PAN|Account|ATM|Password|Address|IFSC|SSN|CVV)\b",
            REDACT,
            text,
            flags=re.IGNORECASE,
        )

        # Names: Title Case (e.g. John Smith, Mary Jane Watson)
        text = re.sub(
            r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2})\b",
            REDACT,
            text,
        )

        # Names: ALL CAPS multi-word (e.g. JOHN DOE, JANE SMITH)
        text = re.sub(
            r"\b([A-Z]{2,}\s+[A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)?)\b",
            REDACT,
            text,
        )

        # Names: Title + Last (e.g. Mr. Smith, Dr. Jane Doe)
        text = re.sub(
            r"\b(Mr|Mrs|Ms|Miss|Dr|Prof)\.?\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b",
            REDACT,
            text,
            flags=re.IGNORECASE,
        )

        # Salary / pay in any currency (rupee, dollar, euro, dinar, pound, etc.)
        text = re.sub(
            r"\b(?:salary|sal|pay|wage|income)\s*[:\-]?\s*\d[\d,.]*\s*(?:lakh|lac|k|cr|million|mn|bn)?\b",
            REDACT,
            text,
            flags=re.IGNORECASE,
        )
        text = re.sub(
            r"[₹$€£]\s*\d[\d,.]*\s*(?:lakh|lac|k|cr|m|mn|bn)?\b",
            REDACT,
            text,
        )
        text = re.sub(
            r"\b(?:rs\.?|inr|rupee|rupees)\s*[:\-]?\s*\d[\d,.]*\s*(?:lakh|lac|k|cr)?\b",
            REDACT,
            text,
            flags=re.IGNORECASE,
        )
        text = re.sub(
            r"\b(?:usd|\$|dollar|dollars)\s*[:\-]?\s*\d[\d,.]*\s*(?:k|m|mn|million|bn)?\b",
            REDACT,
            text,
            flags=re.IGNORECASE,
        )
        text = re.sub(
            r"\b(?:eur|€|euro|euros)\s*[:\-]?\s*\d[\d,.]*\s*(?:k|m|mn|million)?\b",
            REDACT,
            text,
            flags=re.IGNORECASE,
        )
        text = re.sub(
            r"\b(?:gbp|£|pound|pounds)\s*[:\-]?\s*\d[\d,.]*\s*(?:k|m|mn)?\b",
            REDACT,
            text,
            flags=re.IGNORECASE,
        )
        text = re.sub(
            r"\b(?:dinar|dinars|kwd|bhd|jod|iqd|aed|sar)\s*[:\-]?\s*\d[\d,.]*\s*(?:k|m)?\b",
            REDACT,
            text,
            flags=re.IGNORECASE,
        )

        if text != original:
            para.clear()
            para.add_run(text)

    # Redact in tables as well (names, emails, etc.)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text
                    if not t:
                        continue
                    orig = t
                    t = re.sub(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", REDACT, t)
                    t = re.sub(r"\b\d{10}\b", REDACT, t)
                    t = re.sub(r"\b\d{4,}\b", REDACT, t)
                    t = re.sub(
                        r"\b(Aadhaar|PAN|Account|ATM|Password|Address|IFSC)\b",
                        REDACT,
                        t,
                        flags=re.IGNORECASE,
                    )
                    t = re.sub(r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2})\b", REDACT, t)
                    t = re.sub(
                        r"\b([A-Z]{2,}\s+[A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)?)\b",
                        REDACT,
                        t,
                    )
                    # Salary / currency (rupee, dollar, euro, dinar, etc.)
                    t = re.sub(
                        r"\b(?:salary|sal|pay|wage|income)\s*[:\-]?\s*\d[\d,.]*\s*(?:lakh|lac|k|cr|million|mn|bn)?\b",
                        REDACT,
                        t,
                        flags=re.IGNORECASE,
                    )
                    t = re.sub(r"[₹$€£]\s*\d[\d,.]*\s*(?:lakh|lac|k|cr|m|mn|bn)?\b", REDACT, t)
                    t = re.sub(
                        r"\b(?:rs\.?|inr|rupee|rupees)\s*[:\-]?\s*\d[\d,.]*\s*(?:lakh|lac|k|cr)?\b",
                        REDACT,
                        t,
                        flags=re.IGNORECASE,
                    )
                    t = re.sub(
                        r"\b(?:usd|\$|dollar|dollars)\s*[:\-]?\s*\d[\d,.]*\s*(?:k|m|mn|million|bn)?\b",
                        REDACT,
                        t,
                        flags=re.IGNORECASE,
                    )
                    t = re.sub(
                        r"\b(?:eur|€|euro|euros)\s*[:\-]?\s*\d[\d,.]*\s*(?:k|m|mn|million)?\b",
                        REDACT,
                        t,
                        flags=re.IGNORECASE,
                    )
                    t = re.sub(
                        r"\b(?:gbp|£|pound|pounds)\s*[:\-]?\s*\d[\d,.]*\s*(?:k|m|mn)?\b",
                        REDACT,
                        t,
                        flags=re.IGNORECASE,
                    )
                    t = re.sub(
                        r"\b(?:dinar|dinars|kwd|bhd|jod|iqd|aed|sar)\s*[:\-]?\s*\d[\d,.]*\s*(?:k|m)?\b",
                        REDACT,
                        t,
                        flags=re.IGNORECASE,
                    )
                    if t != orig:
                        para.clear()
                        para.add_run(t)

    doc.save(output_path)
